"""
Menyusun Supplementary Table S1 dan S2 untuk INDIC-6889.

Membaca berkas CSV yang dihasilkan multihazard.py:
    validation_metrics_<mode>.csv        -> Tabel S1
    vif_<hazard>_<mode>.csv              -> Tabel S2a
    corr_pearson_<hazard>_<mode>.csv     -> Tabel S2b
    corr_spearman_<hazard>_<mode>.csv    -> Tabel S2b

Keluaran:
    Supplementary_Tables_S1_S2.docx      siap dilampirkan ke naskah
    TableS1_validation.csv
    TableS2a_vif.csv
    TableS2b_correlations.csv

Pakai:
    pip install python-docx pandas
    python make_supplementary.py                 # cari CSV di direktori ini
    python make_supplementary.py --dir /content  # atau di tempat lain
"""

import glob
import os
import re
import sys

import numpy as np
import pandas as pd

# ------------------------------------------------------------------ konfigurasi
# Kalau dipaste ke sel Colab, ubah dua baris ini saja lalu jalankan selnya.
# Kalau dijalankan dari terminal, keduanya bisa ditimpa dengan --dir dan --out.
CSV_DIR = "/content"      # folder berisi CSV keluaran multihazard.py
OUT_DOCX = "Supplementary_Tables_S1_S2.docx"
# Kalau Drive terpasang, salinan hasil ikut disimpan ke sini agar tidak hilang
# saat sesi Colab berakhir. Set None untuk mematikan.
DRIVE_COPY = "/content/drive/MyDrive/GEE_MultiHazard_Revisi"

HAZARD_ORDER = ["landslide", "wildfire", "flood"]
CLF_ORDER = ["RF", "GTB", "CART"]
SCHEME_LABEL = {
    "holdout_70_30": "70:30 hold-out",
    "random_fold": "5-fold random CV",
    "spatial_fold": "5-fold spatial block CV",
}
# Ambang pelaporan pasangan korelasi. 0.5 dipilih agar tabel memuat pasangan
# yang mendekati ambang redundansi 0.7, bukan hanya yang melewatinya, sehingga
# pembaca bisa menilai sendiri.
CORR_REPORT = 0.5
CORR_FLAG = 0.7


# ---------------------------------------------------------------- pemuatan
def find_mode(d):
    """Tentukan absence mode dari nama berkas yang ada."""
    hits = glob.glob(os.path.join(d, "validation_metrics_*.csv"))
    if not hits:
        sys.exit(f"Tidak ada validation_metrics_*.csv di {d}. "
                 "Jalankan multihazard.py lebih dulu.")
    modes = [re.search(r"validation_metrics_(.+)\.csv", os.path.basename(h)).group(1)
             for h in hits]
    if len(modes) > 1:
        print(f"  beberapa mode ditemukan {modes}, memakai '{modes[0]}'")
    return modes[0]


def load_validation(d, mode):
    df = pd.read_csv(os.path.join(d, f"validation_metrics_{mode}.csv"))
    need = {"hazard", "classifier", "scheme", "auc_mean"}
    if not need <= set(df.columns):
        sys.exit(f"kolom kurang di validation_metrics_{mode}.csv: "
                 f"{need - set(df.columns)}")
    return df


def load_vif(d, mode):
    out = {}
    for hz in HAZARD_ORDER:
        p = os.path.join(d, f"vif_{hz}_{mode}.csv")
        if os.path.exists(p):
            out[hz] = pd.read_csv(p)
        else:
            print(f"  lewati VIF {hz}: {os.path.basename(p)} tidak ada")
    return out


def load_corr(d, mode):
    """Kembalikan {hazard: (pearson, spearman)}; spearman None bila belum ada."""
    out = {}
    for hz in HAZARD_ORDER:
        pear = os.path.join(d, f"corr_pearson_{hz}_{mode}.csv")
        spear = os.path.join(d, f"corr_spearman_{hz}_{mode}.csv")
        legacy = os.path.join(d, f"corr_{hz}_{mode}.csv")   # versi lama
        if os.path.exists(pear):
            p = pd.read_csv(pear, index_col=0)
        elif os.path.exists(legacy):
            p = pd.read_csv(legacy, index_col=0)
            print(f"  {hz}: memakai corr_{hz}_{mode}.csv (versi tanpa Spearman)")
        else:
            print(f"  lewati korelasi {hz}: berkas tidak ada")
            continue
        s = pd.read_csv(spear, index_col=0) if os.path.exists(spear) else None
        out[hz] = (p, s)
    return out


# ---------------------------------------------------------------- tabel
def fmt(mean, sd):
    if pd.isna(mean):
        return "\u2014"
    if pd.isna(sd):
        return f"{mean:.3f}"
    return f"{mean:.3f} \u00b1 {sd:.3f}"


def build_s1(val):
    rows = []
    for hz in HAZARD_ORDER:
        for clf in CLF_ORDER:
            sub = val[(val.hazard == hz) & (val.classifier == clf)]
            if sub.empty:
                continue
            row = {"Hazard": hz.capitalize(), "Algorithm": clf}
            for scheme, label in SCHEME_LABEL.items():
                r = sub[sub.scheme == scheme]
                row[label] = (fmt(r.auc_mean.iloc[0], r.auc_sd.iloc[0])
                              if not r.empty else "\u2014")
            rf = sub[sub.scheme == "random_fold"]
            sf = sub[sub.scheme == "spatial_fold"]
            row["\u0394 (spatial \u2212 random)"] = (
                f"{sf.auc_mean.iloc[0] - rf.auc_mean.iloc[0]:+.3f}"
                if not rf.empty and not sf.empty else "\u2014")
            rows.append(row)
    return pd.DataFrame(rows)


def build_s2a(vifs):
    preds = []
    for df in vifs.values():
        preds += [p for p in df.variable if p not in preds]
    rows = []
    for p in preds:
        row = {"Conditioning factor": p}
        for hz in HAZARD_ORDER:
            if hz not in vifs:
                continue
            m = vifs[hz][vifs[hz].variable == p]
            row[hz.capitalize()] = f"{m.VIF.iloc[0]:.2f}" if not m.empty else "\u2014"
        rows.append(row)
    return pd.DataFrame(rows)


def build_s2b(corrs):
    rows = []
    for hz, (pear, spear) in corrs.items():
        cols = list(pear.columns)
        for i, a in enumerate(cols):
            for b in cols[i + 1:]:
                rp = pear.loc[a, b]
                rs = spear.loc[a, b] if spear is not None else np.nan
                if max(abs(rp), abs(rs) if not pd.isna(rs) else 0) < CORR_REPORT:
                    continue
                flag = ("yes" if max(abs(rp), abs(rs) if not pd.isna(rs) else 0)
                        > CORR_FLAG else "")
                lo, hi = sorted((a, b))       # urutan nama dibakukan
                rows.append({
                    "Hazard": hz.capitalize(),
                    "Variable pair": f"{lo} \u2013 {hi}",
                    "Pearson r": f"{rp:+.3f}",
                    "Spearman \u03c1": "\u2014" if pd.isna(rs) else f"{rs:+.3f}",
                    "|r| > 0.7": flag,
                })
    return pd.DataFrame(rows)


# ---------------------------------------------------------------- docx
def add_table(doc, df, style="Light Grid Accent 1"):
    from docx.shared import Pt
    t = doc.add_table(rows=1, cols=len(df.columns))
    try:
        t.style = style
    except KeyError:
        t.style = "Table Grid"
    for j, c in enumerate(df.columns):
        cell = t.rows[0].cells[j]
        cell.text = str(c)
        for r in cell.paragraphs[0].runs:
            r.bold = True
            r.font.size = Pt(9)
    for _, row in df.iterrows():
        cells = t.add_row().cells
        for j, c in enumerate(df.columns):
            cells[j].text = str(row[c])
            for r in cells[j].paragraphs[0].runs:
                r.font.size = Pt(9)
    return t


def write_docx(s1, s2a, s2b, out, has_spearman):
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    for s in doc.sections:
        s.left_margin = s.right_margin = Pt(54)

    doc.add_heading("Supplementary Material", level=1)
    p = doc.add_paragraph()
    p.add_run("Machine Learning-Based Multi-Hazard Adaptation Across Major "
              "Land Systems in Southeast Asia").italic = True
    doc.add_paragraph("Ms. No. INDIC-6889")

    doc.add_heading("Table S1. Model discrimination under three validation "
                    "schemes", level=2)
    doc.add_paragraph(
        "Area under the ROC curve for each algorithm and hazard. The 70:30 "
        "hold-out replicates the original partition. Random cross-validation is "
        "repeated five-fold; spatial block cross-validation assigns whole "
        "one-degree blocks to folds so that no testing observation falls within "
        "the autocorrelation range of a training observation. Values are mean \u00b1 "
        "standard deviation across folds. These results derive from an "
        "independent reimplementation of the workflow and are reported as a "
        "robustness assessment rather than as a reproduction of the primary "
        "results in Section 3.1.")
    add_table(doc, s1)

    doc.add_paragraph()
    doc.add_heading("Table S2a. Variance inflation factors for the conditioning "
                    "factors", level=2)
    doc.add_paragraph(
        "Variance inflation factors computed separately for each hazard on the "
        "continuous conditioning factors. Categorical predictors (land cover, "
        "soil texture class) are excluded because the variance inflation factor "
        "is not meaningful for nominal classes. The design matrix includes a "
        "constant term; without it the statistic is computed from an uncentred "
        "coefficient of determination and is not interpretable.")
    add_table(doc, s2a)

    doc.add_paragraph()
    doc.add_heading("Table S2b. Pairwise correlations among conditioning "
                    "factors", level=2)
    note = (
        "Variable pairs for which either coefficient exceeds 0.5 in absolute "
        "value. Spearman rank correlations are reported alongside Pearson "
        "coefficients because the classifiers used here are tree ensembles, "
        "which are invariant to monotonic transformations of their predictors; "
        "rank correlation is therefore the more appropriate measure of the "
        "redundancy that affects variable importance."
        if has_spearman else
        "Variable pairs for which the Pearson coefficient exceeds 0.5 in "
        "absolute value. Spearman coefficients were not available when this "
        "table was generated."
    )
    doc.add_paragraph(note)
    add_table(doc, s2b)

    doc.save(out)


# ---------------------------------------------------------------- main
def main(csv_dir=None, out=None):
    csv_dir = csv_dir or CSV_DIR
    out = out or OUT_DOCX
    if not os.path.isdir(csv_dir):
        sys.exit(f"folder tidak ada: {csv_dir}  (ubah CSV_DIR di atas)")

    mode = find_mode(csv_dir)
    print(f"  mode: {mode}")

    s1 = build_s1(load_validation(csv_dir, mode))
    s2a = build_s2a(load_vif(csv_dir, mode))
    corrs = load_corr(csv_dir, mode)
    s2b = build_s2b(corrs)
    has_spearman = any(sp is not None for _, sp in corrs.values())

    s1.to_csv("TableS1_validation.csv", index=False)
    s2a.to_csv("TableS2a_vif.csv", index=False)
    s2b.to_csv("TableS2b_correlations.csv", index=False)

    print("\nTabel S1")
    print(s1.to_string(index=False))
    print("\nTabel S2a")
    print(s2a.to_string(index=False))
    print("\nTabel S2b")
    print(s2b.to_string(index=False) if not s2b.empty
          else "  (tidak ada pasangan di atas ambang 0.5)")

    made = ["TableS1_validation.csv", "TableS2a_vif.csv",
            "TableS2b_correlations.csv"]
    try:
        write_docx(s1, s2a, s2b, out, has_spearman)
        made.insert(0, out)
        print(f"\n{out} tersimpan")
    except ImportError:
        print("\npython-docx belum terpasang; CSV tetap tersimpan.")
        print("  pip install python-docx")

    # CSV dan docx ditulis ke direktori kerja, yang di Colab hilang saat sesi
    # berakhir. Salin ke Drive kalau tersedia.
    if DRIVE_COPY and os.path.isdir(os.path.dirname(DRIVE_COPY.rstrip("/"))):
        import shutil
        os.makedirs(DRIVE_COPY, exist_ok=True)
        for f in made:
            if os.path.exists(f):
                shutil.copy(f, DRIVE_COPY)
        print(f"salinan disimpan ke {DRIVE_COPY}")

    return s1, s2a, s2b


# Dijalankan baik saat di-paste ke sel Colab maupun lewat: python make_supplementary.py
if __name__ == "__main__":
    import sys as _sys
    _dir = _sys.argv[_sys.argv.index("--dir") + 1] if "--dir" in _sys.argv else None
    main(_dir)
else:
    main()
